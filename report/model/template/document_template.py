from __future__ import annotations

import os
from pathlib import Path
from docx import Document
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx2pdf import convert
from collections.abc import Callable
from docx.document import _Body
from report.docx.core import add_picture
from report.model.elements.paragraph import Paragraph as MyParagraph
from lxml.etree import _Element

from report.model.element_creator import copy_paragraph
from report.model.elements.formula import Formula
from report.model.template.insert_key import InsertKey
from report.model.template.isolate_key_runs import isolate_key_runs
from report.model.elements.table import Table
from report.model.template.template_util import short_tag, keys_in_run, can_replace_paragraph


def replace_paragraph_with_elements(paragraph: Paragraph, elements: list[_Element]):
    paragraph_element = paragraph._p
    for elem in reversed(elements):
        paragraph_element.addnext(elem)
    body_element = paragraph_element.getparent()
    body_element.remove(paragraph_element)


def get_document_elements(document: DocumentType):
    ignored_tags = ["sectPr"]
    body_element = document._element.body
    result: list[_Element] = []
    for child in body_element.iterchildren():
        if short_tag(child) in ignored_tags:
            continue
        result.append(child)
    return result


def get_image_dimensions(image_path: str):
    from PIL import Image
    # Используем библиотеку PIL для получения размеров изображения
    with Image.open(image_path) as img:
        width, height = img.size
    return width, height


class DocumentTemplate:
    def __init__(self, path: Path):
        self.document: DocumentType = Document(path.as_posix())
        self.name = path.name
        self._insert_keys = None
        self._last_key_index = 0

    def replace_with_template(self, key: str, document_template: DocumentTemplate):
        other_iks = list(document_template.insert_keys.values())
        self.insert_data_from_document(key, document_template.document)

        start_i = self._last_key_index
        # other_iks = list(document_template.insert_keys.values())
        end_i = len(other_iks) + start_i

        new_body = _Body(self.document.element.body, parent=self.document)

        for new_ik_index, old_ik in zip(range(start_i, end_i), other_iks):
            para: Paragraph = old_ik.paragraph
            para._parent = new_body
            self._insert_keys[new_ik_index] = old_ik
        self._last_key_index = end_i

    def save(self, save_path: Path, document_name: str = "output.docx", add_pdf: bool = True):
        docx_path = save_path.joinpath(document_name)
        self.document.save(docx_path.absolute().as_posix())
        if add_pdf:
            pdf_path = save_path.joinpath(f"{document_name.split('.')[0]}.pdf")
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
            convert(docx_path.absolute().as_posix(), pdf_path.absolute().as_posix())

    @property
    def insert_keys(self) -> dict[int, InsertKey]:
        """
            Парсинг ключей делается с предположением, что в ключи находятся в обособленных Run, т.к. были
            выполнены преобразования в init
         """
        if self._insert_keys is not None:
            return self._insert_keys

        isolate_key_runs(self.document)

        result = []
        paragraphs: list[Paragraph] = self.document.paragraphs
        for p in paragraphs:
            for r in p.runs:
                keys = keys_in_run(r)
                if keys:
                    insert_key = InsertKey(
                        paragraph=p,
                        run=r,
                        key=keys[0].strip()
                    )
                    result.append(insert_key)
        dict_result = {}
        for i, key in enumerate(result):
            dict_result[i] = key
        self._insert_keys = dict_result
        self._last_key_index = len(result)
        return dict_result

    def _find_insert_keys(self, key: str) -> dict[int, InsertKey]:
        result = {}
        for i, ik in self.insert_keys.items():
            if ik.key == key:
                result[i] = ik
        if len(result) == 0:
            print(f"Ключ [{key}] не существует или уже использован")
        return result

    def _insert(self, key: str, insert_func: Callable[[InsertKey], None], report: str):
        insert_keys = self._find_insert_keys(key)
        for i, ik in insert_keys.items():
            insert_func(ik)
            del self.insert_keys[i]
            print(f"{self.name}:[{key}]: {report}")

    def insert_text(self, key: str, text: str):
        def insert_func(ik: InsertKey):
            ik.run.text = text

        self._insert(key=key, insert_func=insert_func, report="TEXT")

    def insert_formula(self, key: str, formula: Formula):
        def insert_func(ik: InsertKey):
            ik.run._r.addnext(formula.oMath)
            para = ik.run._r.getparent()
            para.remove(ik.run._r)

        self._insert(key=key, insert_func=insert_func, report="FORMULA")

    def insert_data_from_document(self, key: str, document: DocumentType):
        def insert_func(ik: InsertKey):
            if not can_replace_paragraph(ik.paragraph):
                print(
                    f"Нельзя вставить содержимое другого документа по ключу [{key}] вместо параграфа [{ik.paragraph.text}],"
                    f"т.к в параграфе есть другой текст помимо ключа")
                return
            elements_to_insert = get_document_elements(document)
            replace_paragraph_with_elements(ik.paragraph, elements_to_insert)

        self._insert(key=key, insert_func=insert_func, report="DOCUMENT")

    def insert_table(self, key: str, table: Table):
        def insert_func(ik: InsertKey):
            if not can_replace_paragraph(ik.paragraph):
                print(
                    f"Нельзя вставить таблицу по ключу [{key}] вместо параграфа [{ik.paragraph.text}],"
                    f"т.к в параграфе есть другой текст помимо ключа")
                return

            table_element = table.table_element
            replace_paragraph_with_elements(ik.paragraph, [table_element])

        self._insert(key=key, insert_func=insert_func, report="TABLE")

    def insert_picture(self, key: str, picture_path: Path):
        def insert_func(ik: InsertKey):
            add_picture(picture_path.as_posix(), ik.run)

        self._insert(key=key, insert_func=insert_func, report="PICTURE")

    def insert_elements_list(self, key: str, elements_list: list[Table | Formula | MyParagraph | DocumentType]):
        def insert_func(ik: InsertKey):
            if not can_replace_paragraph(ik.paragraph):
                print(
                    f"Нельзя вставить список элементов по ключу [{key}] вместо параграфа [{ik.paragraph.text}],"
                    f"т.к в параграфе есть другой текст помимо ключа")
                return

            elements_to_insert = []
            for elem in elements_list:
                if isinstance(elem, Table):
                    elements_to_insert.append(elem.table_element)
                elif isinstance(elem, Formula):
                    p = copy_paragraph(ik.paragraph)
                    p.text = ""
                    p_elem = p._p
                    p_elem.append(elem.oMath)
                    elements_to_insert.append(p_elem)
                elif isinstance(elem, MyParagraph):
                    p = elem.element
                    elements_to_insert.append(p)
                elif isinstance(elem, DocumentType):
                    elements_to_insert.extend(get_document_elements(elem))
                else:
                    raise ValueError(f"Неподдерживаемый тип вставки элемента {type(elem)}")
            replace_paragraph_with_elements(ik.paragraph, elements_to_insert)

        self._insert(key=key, insert_func=insert_func, report="ELEMENTS_LIST")

    def delete_key(self, key: str):
        def insert_func(ik: InsertKey):
            if can_replace_paragraph(ik.paragraph):
                parent = ik.paragraph._p.getparent()
                parent.remove(ik.paragraph._p)
            else:
                parent = ik.paragraph._p
                parent.remove(ik.run._r)

        self._insert(key=key, insert_func=insert_func, report=" -- DELETED -- ")
