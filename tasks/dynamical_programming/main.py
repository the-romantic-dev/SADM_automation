from itertools import combinations

from docx import Document
from docx.document import Document as DocumentType
from igraph import Graph

from report.docx.omml import latex2omml
from report.docx.tables import create_table_filled
from report.docx.templates import fill_template
from solvers.graph_solver import GraphSolver
from tasks.dynamical_programming.fillers.graph.graph_filler import GraphDocxFiller
from tasks.dynamical_programming.solvers.resource_allocation_solver import ResourceAllocationSolver
from tasks.dynamical_programming.solvers.troubleshooting_solver import TroubleshootingSolver


def fill_graph_docx():
    edges = [
        # Горизонтальные ребра
        (1, 2, 2), (1, 3, 5), (1, 4, 7),
        (2, 5, 8),
        (3, 5, 3), (3, 6, 1),
        (4, 6, 3), (4, 7, 6), (4, 8, 1), (4, 9, 4),
        (5, 11, 6), (5, 9, 10),
        (6, 10, 6),
        (7, 12, 3),
        (8, 7, 8), (8, 9, 4), (8, 12, 5),
        (9, 10, 1), (9, 11, 1),
        (10, 12, 2),
        (11, 10, 5), (11, 12, 4)
    ]
    graph_solver = GraphSolver(edges)
    graph_solver.show_graph()
    graph_data_producer = GraphDocxFiller(graph_solver, 1, 12).get_data_producers()
    template: DocumentType = Document("document_template/graph/graph_template.docx")
    fill_template(template=template, data_producers=graph_data_producer)
    template.save("fillers/graph/graph.docx")


def fill_troubleshooting_docx():
    st = [
        [False, False, True, True],
        [False, True, True, True],
        [True, False, False, True],
        [False, False, False, True],
        [True, True, True, False]
    ]
    p = [0.15, 0.15, 0.2, 0.3, 0.2]
    c = [10, 15, 15, 10]
    troubleshooting_solver = TroubleshootingSolver(st=st, p=p, c=c)
    combs = troubleshooting_solver.get_optimal_average_losses_combinations_data(5)

    doc_with_table: DocumentType = Document()
    table_header = [
        latex2omml("R"), latex2omml(r"\bar{f}\left( R \right)"), "Оптимальный тест"
    ]
    table_data = [table_header] + list(map(lambda elem: [";".join([str(i) for i in elem[0]]), elem[1], elem[2]], combs))
    graph: Graph = Graph(directed=True)
    troubleshooting_solver.get_test_graph(graph=graph, r=[1, 2, 3, 4, 5], combs=combs)
    gplot, _ = troubleshooting_solver._prepare_plot(graph)
    gplot.show()
    table = create_table_filled(document=doc_with_table, data=table_data)
    table.autofit = True
    doc_with_table.save("./fillers/troubleshooting/troubleshooting.docx")


def fill_resource_allocation_docx():
    c = [1, 2, 3, 4, 5]
    g = [
        [0.2, 0.8, 1.1, 2.3, 2.9],
        [0.5, 0.9, 1.3, 1.7, 2.7],
        [0.3, 0.6, 1.4, 1.9, 2.6],
        [0.6, 0.9, 1.6, 1.8, 2.4]
    ]
    resource_allocation_solver = ResourceAllocationSolver(c=c, g=g)
    resource_allocation_solver.solve()
    doc: DocumentType = Document()
    calculation_tables = resource_allocation_solver.calculation_tables
    x_columns = resource_allocation_solver.x_columns
    f_columns = resource_allocation_solver.f_columns
    max_row_len = -1
    for row in calculation_tables:
        if len(row) > max_row_len:
            max_row_len = len(row)

    for i in range(len(calculation_tables)):
        data = []
        table = calculation_tables[i]
        for row in table:
            data.append([])
            last_row_index = len(data) - 1
            for j in range(max_row_len + 1):
                value = ""
                if j < len(row):
                    value = f"{row[j][0]}"
                data[last_row_index].append(value)
        for j in range(len(x_columns[i])):
            data[j].append(x_columns[i][j])

        for j in range(len(f_columns[i])):
            data[j].append(f_columns[i][j])
        print()
        doc.add_paragraph()
        create_table_filled(document=doc, data=data)

    doc.save("./fillers/resource_allocation/resource_allocation.docx")


if __name__ == '__main__':
    fill_graph_docx()
    # fill_troubleshooting_docx()
    # fill_resource_allocation_docx()
