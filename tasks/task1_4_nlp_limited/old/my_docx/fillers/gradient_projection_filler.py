from typing import List

from docx import Document
from docx.text.paragraph import Paragraph
from sympy import latex, Matrix
from docx.document import Document as DocumentType
from methods.gradient_projection import GradientProjection
from my_docx.docx_output import sympy2omml, latex2omml, fill_paragraph

from report.model.docx_parts.formula import Formula
from report.model.elements.math.braces import braces, BraceType
from report.model.elements.math.matrix import matrix_from_sympy


class GradientProjectionDocxFiller:
    def __init__(self, gradient_projection):
        self.gp: GradientProjection = gradient_projection
        self.step_paragraphs_producer = lambda: Document(
            'report_data/gradient_projections_steps_temple.docx').paragraphs

    def _create_gp_A0(self):
        return sympy2omml(self.gp.A0)

    def _create_gp_b(self):
        return sympy2omml(self.gp.b)

    def _create_gp_steps(self):
        return self._generate_steps()

    def _create_gp_pic(self):
        return "report_data/gradient_projection.png"

    def _crate_gp_X(self):
        return latex2omml(
            'X^* = ' +
            latex(Matrix([self.gp.solution['x1'], self.gp.solution['x2']])) +
            r'\approx' +
            latex(Matrix([round(float(self.gp.solution['x1']), 4), round(float(self.gp.solution['x2']), 4)])))

    def _create_gp_f(self):
        return latex2omml(
            'f = ' +
            latex(self.gp.solution['f']) +
            r'\approx' +
            latex(round(float(self.gp.solution['f']), 4)))

    def get_data_producers(self):
        return {
            'gp_A0': self._create_gp_A0,
            'gp_b': self._create_gp_b,
            'gp_steps': self._create_gp_steps,
            'gp_pic': self._create_gp_pic,
            'gp_X': self._crate_gp_X,
            'gp_f': self._create_gp_f
        }

    def _create_I_pred(self, data):
        result = list(data['t_пред_i'].keys())
        result.remove(-1)
        return latex2omml(' , '.join([f'{i + 1}' for i in result]))

    def _create_t_chooser(self, data, solution_index):
        result = list(data['t_пред_i'].keys())
        result.remove(-1)
        return latex2omml('t^{(' + str(solution_index) + r')}=\min(t^*,{' + ','.join(
            ['t_{пред_' + str(i + 1) + '}' for i in result]) + '})')

    def _create_t_default(self, data, solution_index):
        return latex2omml(
            r'-\frac{∇^T f(X^{(i)} )⋅K^{(i)}}{{K^{(i)}}^T⋅H(X^{(i)})⋅K^{(i)}}='
            .replace('i', str(solution_index)) + latex(data['t*']))

    def _create_t_preds(self, data, solution_index):
        doc: DocumentType = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        preds = list(data['t_пред_i'].keys())
        preds.remove(-1)
        for n in preds:
            run.element.append(latex2omml(
                r't_{пред_n}=\frac{b_n-a_n X^{(i)}}{a_n K^{(i)}} = '.
                replace('n', str(n + 1)).
                replace('i', str(solution_index)) + latex(data['t_пред_i'][n])
            ))
            run.add_break()
        return run

    def _create_t_result(self, data, solution_step):
        return latex2omml('t^{(i)} = '.replace('i', str(solution_step)) + latex(data['result_step_size']))

    def _create_next_solution(self, data, solution_index):
        return latex2omml('X^{(next)}=X^{(curr)}+t^{(curr)}K^{(curr)}='
                          .replace('next', str(solution_index + 1))
                          .replace('curr', str(solution_index)) +
                          latex(data['Xi+1']))

    def _create_step_2_label(self, data, solution_index):
        # text1 = ["Так как предыдущее направление было ", "{{t_pred}}", ", то мы находимся на границе ", "{{limit}}"]
        text1 = ['Найдем текущие активные ограничения: ', "{{active_limits_formula}}",
                 ". Видим, что активные ограничения это ", "{{active_limits}}"]
        text2 = ["В точке ", "{{current_solution}}",
                 " матрица активных ограничений A и оператор проекции P остаются прежними"]
        doc = Document()
        paragraph = doc.add_paragraph()

        # if data['limit_index'] != -1:
        if len(data['limit_index']) > 0:
            for part in text1:
                run = paragraph.add_run()
                run.text = part
            active_limits_formula = Formula([
                f"A_0 \\cdot X^{{({solution_index})}} - b = ",
                matrix_from_sympy(data['limits_matrix'], brace_type=BraceType.BRACKETS)
            ])
            fill_paragraph(paragraph, {
                # '{{active_limits_formula}}': lambda: latex2omml("t_{пред_i}".replace('i', str(data['limit_index'] + 1))),
                'active_limits_formula': lambda: active_limits_formula.oMath,
                # 'active_limits': lambda: latex2omml("a_i".replace('i', str(data['limit_index'] + 1)))
                'active_limits': lambda: Formula(','.join([f'a_{i + 1}' for i in data['limit_index']])).oMath
            })
        else:
            for part in text2:
                run = paragraph.add_run()
                run.text = part
            fill_paragraph(paragraph, {
                'current_solution': lambda: latex2omml('X^{(i)}'.replace('i', str(solution_index)))
            })
        return paragraph

    def _create_newA(self, data):
        result = latex2omml("A = " + latex(data["newA"]))
        return result

    def _create_new_grad(self, data, solution_index):
        result = latex2omml(
            "∇f(X^{(i)})=".
            replace('i', str(solution_index)) +
            latex(data['new_grad']))
        return result

    def _create_K(self, data, solution_index):
        return latex2omml(
            "K^{(i)} = P∇f(X^{(i)}) = ".replace('i', str(solution_index)) +
            latex(data['K'])
        )

    def _create_lamb_direction(self, data, solution_index):
        # return latex2omml(
        #     (r"K^{(i)} = \arg \underset{\underset{\lambda_k>0}{k}}{\max}\left\{" +
        #     r" \left\| \widetilde{P}_k\nabla f(X^{(i)}) \right\| \right\}=").
        #     replace("(i)", f"({solution_index})") + latex(data["lamb_direction"])
        # )
        i = solution_index
        formula_data = [
            f'K^{{({i})}} = ',
            '\\underset{\\underset{\\lambda_k>0}{k}}{argmax}',
            braces(
                braces(
                    latex2omml(f'\\widetilde{{P}}_k\\nabla f(X^{{({i})}})'),
                    brace_type=BraceType.DOUBLE_STRAIGHT
                ),
                brace_type=BraceType.CURLY
            ),
            '=',
            matrix_from_sympy(data["lamb_direction"], brace_type=BraceType.BRACKETS)
        ]
        return Formula(formula_data).oMath

    def _create_gradient_direction(self, data, solution_index):
        return latex2omml(r"K^{(i)}=\nabla f(X^{(i)})=".
                          replace("(i)", f'({solution_index})') + latex(data["gradient_direction"]))

    def _generate_steps(self):
        from ..docx_output import fill_paragraph

        temp: DocumentType = Document()
        temp_paragraphs = []
        solution_index = 0
        for data in self.gp.report_data:
            step = data['step']
            step_paragraph = self.step_paragraphs_producer()[step]
            data_producers = {}

            match data['step']:
                case 0:
                    data_producers = {
                        'A*grad': lambda: sympy2omml(data['A'] * data['gradient']),
                        'start_direction': lambda: sympy2omml(data['gradient'])
                    }
                case 1:
                    data_producers = {
                        "A0*K": lambda: latex2omml(
                            'K^{(i)} = '.replace('i', str(solution_index)) + latex(data['limitation_vector'])),
                        "I_пред": lambda: self._create_I_pred(data),
                        "t_chooser": lambda: self._create_t_chooser(data, solution_index),
                        "t*": lambda: self._create_t_default(data, solution_index),
                        't_preds': lambda: self._create_t_preds(data, solution_index),
                        't_result': lambda: self._create_t_result(data, solution_index),
                        'next_solution': lambda: self._create_next_solution(data, solution_index)
                    }

                case 2:
                    data_producers = {
                        "step_2_label": lambda: self._create_step_2_label(data, solution_index),
                        "newA": lambda: self._create_newA(data),
                        "new_grad": lambda: self._create_new_grad(data, solution_index),
                        "P": lambda: sympy2omml(data['P']),
                        "K": lambda: self._create_K(data, solution_index)
                    }
                case 3:
                    data_producers = {
                        "lamb": lambda: sympy2omml(data["lamb"])
                    }
                case 4:
                    data_producers = {
                        "lamb": lambda: sympy2omml(data["lamb"]),
                        "lamb_direction": lambda: self._create_lamb_direction(data, solution_index)
                    }
            fill_paragraph(step_paragraph, data_producers)
            if data['step'] == 1:
                solution_index += 1
            # temp.element.append(step_paragraph._element) # Не добавляет новый параграф
            # p = temp.add_paragraph()
            # p._element.append(step_paragraph._element) # Не открывается документ
            temp_paragraphs.append(step_paragraph)

        start_p = temp.add_paragraph()
        for p in reversed(temp_paragraphs):
            start_p._p.addnext(p._element)
        return temp
