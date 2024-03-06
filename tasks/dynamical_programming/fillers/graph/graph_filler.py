from typing import Dict, Callable
from docx import Document
from docx.document import Document as DocumentType
from report.DocxFiller import DocxFiller
from report.docx.omml import latex2omml

from tasks.dynamical_programming.solvers.graph_solver import GraphSolver


class GraphDocxFiller(DocxFiller):
    def __init__(self, graph_solver: GraphSolver, start, end):
        self.graph_solver = graph_solver
        self.start = start
        self.end = end
        self.levels = self.graph_solver.get_vertices_levels()
        self.min_path_size, self.min_path_steps, self.min_path = self.graph_solver.get_min_path(self.start, self.end)
        self.max_path_size, self.max_path_steps, self.max_path = self.graph_solver.get_max_path(self.start, self.end)

    def get_data_producers(self) -> Dict[str, Callable]:
        return {
            "shortest_steps": lambda: self._produce_solution_steps(arg="min"),
            "graph_levels": self._produce_graph_levels,
            "shortest_path": lambda: latex2omml(" → ".join([f"{i}" for i in self.min_path])),
            "shortest_size": lambda: latex2omml(f"{self.min_path_size}"),
            "longest_steps": lambda: self._produce_solution_steps(arg="max"),
            "longest_path": lambda: latex2omml(" → ".join([f"{i}" for i in self.max_path])),
            "longest_size": lambda: latex2omml(f"{self.max_path_size}")
        }

    def _produce_graph_levels(self):
        result_latex = ""
        levels = self.levels
        for l in levels:
            elements = ""
            for e in l:
                if l.index(e) != len(l) - 1:
                    elements += f"{e}, "
                else:
                    elements += f"{e}"
            if levels.index(l) != len(levels):
                result_latex += f"\\left( {elements} \\right), "
            else:
                result_latex += f"\\left( {elements} \\right)"
        return latex2omml(result_latex)

    def _produce_graph_img(self):
        path = "graph.png"
        self.graph_solver.save_graph(path)
        return path

    def _produce_solution_steps(self, arg: str):
        q_minmax_latex = lambda start, end: f"q^{{{arg}}}\left({start}, {end} \\right)"
        if arg == "max":
            path_steps = self.max_path_steps
        elif arg == "min":
            path_steps = self.min_path_steps
        else:
            raise ValueError("Тип пути не min и не max")

        def get_step_by_start(start) -> Dict:
            for _step in path_steps:
                if _step['start'] == start:
                    return _step

        def get_variants_latex(_step: dict) -> str:
            result = ""
            for out in _step['outs']:
                _latex = f"{q_minmax_latex(out, _step['end'])} + q \\left({_step['start']}, {out} \\right)"
                result += f"{_latex} \\\\"
            return result

        temp_doc: DocumentType = Document()
        step_counter = 0
        for level in reversed(self.levels):
            temp_doc.add_paragraph(f"Шаг {step_counter + 1}")
            step_counter += 1
            for element in level:
                p = temp_doc.add_paragraph()
                r = p.add_run()

                if element == self.end:
                    res_latex = f"{q_minmax_latex(element, element)} = 0"
                else:
                    step = get_step_by_start(element)
                    variants_latex = get_variants_latex(step)
                    variants_values_latex = " \\\\".join([f"{i}" for i in step['path_sizes']])
                    res_latex = (f"{q_minmax_latex(step['start'], step['end'])} ="
                                 f"\\{arg}\\left( \\begin{{matrix}} {variants_latex} \\end{{matrix}} \\right) ="
                                 f"\\{arg}\\left( \\begin{{matrix}} {variants_values_latex} \\end{{matrix}} \\right) ="
                                 f"{step['opt_path_size']}")
                r.element.append(latex2omml(res_latex))

        return temp_doc
