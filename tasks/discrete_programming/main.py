from report.docx.templates import fill_paragraph_template, fill_template
from tasks.discrete_programming.tsp.filler import TSPStepDocxFiller, TSPLastStepDocxFiller, TSPZeroStepDocxFiller
from tasks.discrete_programming.tsp.solver import TSPSolver
from docx.document import Document as DocumentType
from docx import Document

if __name__ == "__main__":
    tsp_solver = TSPSolver(list_matrix=[
        ["inf", 31, 26, 18, 20, 32],
        [22, "inf", 22, 48, 16, 27],
        [20, 20, "inf", 30, 6, 6],
        [17, 35, 22, "inf", 36, 25],
        [19, 14, 6, 25, "inf", 12],
        [23, 25, 6, 21, 12, "inf"],
    ])
    report_data = []
    tsp_solver.solve_min(report_data=report_data)

    result_doc: DocumentType = Document()
    last_step = len(report_data) - 1
    for step, data in enumerate(report_data):
        if data.tree_level == -1:
            tsp_filler = TSPZeroStepDocxFiller(step_report_data=data)
            step_template: DocumentType = Document("tsp/step_template_zero.docx")
        elif step != last_step:
            tsp_filler = TSPStepDocxFiller(step_report_data=data, step_number=step)
            step_template: DocumentType = Document("tsp/step_template.docx")
        else:
            tsp_filler = TSPLastStepDocxFiller(step_report_data=data, step_number=step)
            step_template: DocumentType = Document("tsp/step_template_last.docx")
        paragraph = result_doc.add_paragraph("{{step}}")
        fill_template(template=step_template, data_producers=tsp_filler.get_data_producers())
        fill_paragraph_template(paragraph, data_producers={"step": lambda: step_template})

    result_doc.save("tsp.docx")
