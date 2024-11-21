from docx import Document
from docx.document import Document as DocumentType
from sympy import symbols, latex

import util
from my_docx.docx_output import latex2omml
from task import Task

class OptimumConditionDocxFiller:

    def get_data_producers(self):
        return {
            'linear_limits_kun_takker': lambda: self._create_linear_limits_kun_takker(Task.lim1234),
            'quadratic_limits_kun_takker': lambda: self._create_linear_limits_kun_takker(Task.lim67)
        }

    def _create_linear_limits_kun_takker(self, limits):
        doc: DocumentType = Document()
        main_function_takker = util.gradient()
        for i in range(len(limits)):
            u = symbols(f'u{i+1}')
            main_function_takker += u * util.gradient(expr=limits[i])
        main_function_takker = main_function_takker.T.tolist()[0]
        limits_takker = []
        for i in range(len(limits)):
            u = symbols(f'u{i+1}')
            limits_takker.append(u * limits[i])

        result_takker = main_function_takker + limits_takker
        for i in range(len(result_takker)):
            doc.add_paragraph().add_run().element.append(
                latex2omml(latex(result_takker[i]) + " = 0")
            )
        return doc