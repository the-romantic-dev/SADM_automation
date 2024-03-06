import re

from docx.shared import Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from sympy import Rational, Symbol, pretty, latex, Add, Expr
from docx.document import Document as DocumentType
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
from typing import List, Dict, Any
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import latex2mathml.converter
from lxml import etree
from lxml.etree import _XSLTResultTree, _Element

from report.docx.core import get_image_dimensions

document_name = 'output.docx'


def paint_table_row(table, row_index, color='FFE699'):
    for _, cell in enumerate(table.rows[row_index].cells):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)


def paint_table_column(table, column_index, color='FFE699'):
    for _, cell in enumerate(table.columns[column_index].cells):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)


def paint_table_cell(table, row_index, column_index, color='FFE699'):
    column_size = len(table.columns)
    for i, cell in enumerate(table.rows[row_index].cells):
        if 0 <= column_index == i or (column_index < 0 and i == column_size + column_index):
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)
            break


def fill_table(table, data: List[List], alignment='center'):
    if len(table.rows) != len(data):
        raise IndexError("Количество строк в data не совпадает с количеством строк в таблице")
    for i, row in enumerate(data):
        for j, cell_value in enumerate(row):
            cell = table.cell(i, j)
            run = cell.paragraphs[0].add_run()
            if isinstance(cell_value, Rational) or isinstance(cell_value, Symbol) or isinstance(cell_value, Expr):
                # input_rational_for_docx(run=cell.paragraphs[0].runs[0], p=cell_value.p, q=cell_value.q)
                run.element.append(sympy2omml(cell_value))
            elif isinstance(cell_value, str):
                run.element.append(latex2omml(cell_value))
            else:
                cell.text = str(cell_value)
            match alignment:
                case 'center':
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                case 'left':
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                case 'right':
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def create_table(document: DocumentType, rows: int, cols: int):
    result = document.add_table(rows, cols)
    result.style = 'Table Grid'
    return result


def create_table_filled(
        document: DocumentType,
        data: List[List],
        alignment='center'
):
    result = create_table(document, len(data), len(data[0]))
    fill_table(table=result, data=data, alignment=alignment)
    return result


def input_rational_for_docx(run, p, q):
    oMath = OxmlElement('m:oMath')
    frac = OxmlElement('m:f')
    num = OxmlElement('m:num')
    den = OxmlElement('m:den')

    num.append(OxmlElement('m:r'))
    num[0].append(OxmlElement('m:t'))
    num[0][0].text = f'{p}'

    den.append(OxmlElement('m:r'))
    den[0].append(OxmlElement('m:t'))
    den[0][0].text = f'{q}'

    frac.append(num)
    frac.append(den)

    mr = OxmlElement('m:r')
    mr.append(frac)
    oMath.append(mr)

    run._element.clear_content()
    # Set the font size for the equation

    # Add the math element to the run
    run._element.append(oMath)


def latex2omml(latex_expr):
    mathml_output = latex2mathml.converter.convert(latex=latex_expr)
    mml2omml_path = r'D:\Убежище\Университет\5 семестр\СисАнал\Скрипты\task_4\my_docx\MML2OMML.XSL'

    tree = etree.fromstring(mathml_output)
    xslt = etree.parse(mml2omml_path)
    transform = etree.XSLT(xslt)
    new_dom = transform(tree)
    return new_dom.getroot()


def sympy2omml(sympy_expr):
    latex_output = latex(sympy_expr)
    return latex2omml(latex_output)


def contains_template(string):
    pattern = r'\{\{[^\}]+\}\}'  # Регулярное выражение для поиска шаблона {{...}}
    match = re.search(pattern, string)
    return bool(match)


def extract_name_from_template(string):
    pattern = r'\{\{([^}]+)\}\}'  # Обновленное регулярное выражение с захватывающими скобками
    match = re.search(pattern, string)
    if match:
        return match.group(1)  # Извлекаем значение из захваченной группы
    else:
        return None


def fill_paragraph(paragraph: Paragraph, data_producers: Dict[str, Any]):
    pieced_name = ""
    is_piecing_started = False
    for run in paragraph.runs:
        run: Run = run
        name = extract_name_from_template(run.text)
        if name is None:
            if "{{" in run.text:
                is_piecing_started = True
                pieced_name += run.text.split('{{')[1]
                run.text = ''
            elif "}}" in run.text:
                is_piecing_started = False
                pieced_name += run.text.split('}}')[0]
                name = pieced_name
                run.text = "{{" + pieced_name + "}}"
                pieced_name = ''
            else:
                if is_piecing_started:
                    pieced_name += run.text
                    run.text = ''
                if not is_piecing_started:
                    # print(f'В run нет ключа. Текст run: [{run.text}]')
                    pass
        if name in data_producers.keys():
            process_run_append(run=run, data_producer=data_producers[name])
        elif name is not None:
            print(f'Во входном словаре для заполнения нет ключа [{name}]')


def add_picture(picture_path: str, run: Run):
    run.element.text = ''
    pic = run.add_picture(picture_path)
    original_width, original_height = get_image_dimensions(picture_path)
    pic.width = Pt(612)
    pic.height = Pt(int((612 / original_width) * original_height))


def add_text(text: str, run: Run):
    run.element.text = ''
    # run.text = run.element.xml.replace('{{' + keyname + '}}', str(current_data))
    run.element.text = text


def _make_table_from_child_document(child, document):
    table = Table(child, document)
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={'sz': 2, 'val': 'single', 'color': '#000000'},
                bottom={'sz': 2, 'val': 'single', 'color': '#000000'},
                start={'sz': 2, 'val': 'single', 'color': '#000000'},
                end={'sz': 2, 'val': 'single', 'color': '#000000'}
            )
    return table


def add_document_content(run: Run, document: DocumentType):
    from docx.oxml import CT_P
    from docx.oxml import CT_Tbl

    run.element.text = ''
    parent_elm = document.element.body
    paste_elements = []
    for child in parent_elm.iterchildren():
        result = None
        if isinstance(child, CT_P):
            result = Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            result = _make_table_from_child_document(child, document)
        else:
            continue
        paste_elements.append(result._element)
    for elem in reversed(paste_elements):
        run._r.addnext(elem)


def process_run_append(run, data_producer):
    current_data = data_producer()
    if isinstance(current_data, str):
        " or isinstance(current_data, _XSLTResultTree)"
        if current_data.endswith('.png'):
            print(f'Вставлено изображение {run.text}')
            add_picture(picture_path=current_data, run=run)
        else:
            print(f'Вставлен текст {run.text}')
            add_text(text=current_data, run=run)
    elif isinstance(current_data, _Element):
        print(f'Вставлен элемент {run.text}')
        run.element.text = ''
        run.element.append(current_data)
    elif isinstance(current_data, DocumentType):
        print(f'Обработана вставка из документа в {run.text}')
        add_document_content(run=run, document=current_data)
    elif isinstance(current_data, Run):
        print(f'Вставка Run в {run.text}')
        run.element.text = ''
        run._r.append(current_data.element)
    elif isinstance(current_data, Paragraph):
        print(f'Вставка Paragraph в {run.text}')
        run.text = ''
        for r in current_data.runs:
            run._r.append(r.element)


    else:
        print(f'Неподдерживаемый для вставки тип значения {type(current_data)}')


def fill_template(template: DocumentType, data_producers: Dict[str, Any]):
    for i, paragraph in enumerate(template.paragraphs):
        paragraph: Paragraph = paragraph
        if contains_template(paragraph.text):
            fill_paragraph(paragraph, data_producers)


def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
