from sympy import Rational, Symbol,pretty
from docx.document import Document as DocumentType
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml, OxmlElement
from typing import List
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# document_name = 'output.docx'

def paint_table_row(table, row_index, color ='FFE699'):
    for _, cell in enumerate(table.rows[row_index].cells):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

def paint_table_column(table, column_index, color ='FFE699'):
    for _, cell in enumerate(table.columns[column_index].cells):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

def paint_table_cell(table, row_index, column_index, color ='FFE699'):
    for i, cell in enumerate(table.rows[row_index].cells):
        if i == column_index:
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)
            break

def fill_table(table, data: List[List], alignment = 'center'):
    if len(table.rows) != len(data):
        raise IndexError("Количество строк в util не совпадает с количеством строк в таблице")
    for i, row in enumerate(data):
        for j, cell_value in enumerate(row):
            cell = table.cell(i, j)
            cell.paragraphs[0].add_run()
            if  isinstance(cell_value, Rational) and cell_value.q != 1:
                input_rational_for_docx(run=cell.paragraphs[0].runs[0],p=cell_value.p, q=cell_value.q)
            elif isinstance(cell_value, Symbol):
                cell.text = pretty(cell_value)
            else:
                cell.text = str(cell_value)
            match alignment:
                case 'center':
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                case 'left':
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                case 'right':
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

def create_table(document: DocumentType, rows: int, cols: int, header: str = "\n"):
    document.add_paragraph(header)
    result = document.add_table(rows, cols)
    result.style = 'Table Grid'
    return result

def create_table_filled(
    document: DocumentType,
    data: List[List],
    header: str = "\n",
    alignment = 'center'
    ):
    result = create_table(document, len(data), len(data[0]), header=header)
    fill_table(table=result, data=data, alignment=alignment)
    return result

def input_rational_for_docx(run, p, q):
    oMath = OxmlElement('m:oMath')
    frac = OxmlElement('m:f')
    num = OxmlElement('m:num')
    den = OxmlElement('m:den')

    num.append(OxmlElement('m:r'))
    num[0].append(OxmlElement('m:t'))
    num[0][0].text = f'{p}'

    den.append(OxmlElement('m:r'))
    den[0].append(OxmlElement('m:t'))
    den[0][0].text = f'{q}'

    frac.append(num)
    frac.append(den)

    mr = OxmlElement('m:r')
    mr.append(frac)
    oMath.append(mr)

    run._element.clear_content()
    # Set the font size for the equation

    # Add the math element to the run
    run._element.append(oMath)