from typing import Dict, Callable
from docx.document import Document as DocumentType
from docx import Document
from report.DocxFiller import DocxFiller
from report.docx.omml import latex2omml
from report.docx.tables import create_table_filled, paint_table_row, paint_table_column
from .report import TSPReportDataMinimal, TSPReportData, TSPReportDataLast
from .tsp_city_matrix import TSPCityMatrix


def paths_way_to_latex(paths_way):
    result = []
    for path in paths_way:
        if path[1]:
            result.append(str((path[0][0] + 1, path[0][1] + 1)))
        else:
            result.append(f"\\left( \\overline{{ {path[0][0] + 1}, {path[0][1] + 1} }} \\right)")
    str_res = "()"
    if len(result) > 0:
        str_res = ",".join(result)
    return str_res


def get_exclude_G_latex(report_data: TSPReportData) -> str:
    paths_way = report_data.paths_way
    worst_tau = report_data.worst_tau
    data = paths_way + [(worst_tau[0], False)]
    return f"G_{{ {paths_way_to_latex(data)} }}^{{ ({report_data.tree_level + 1})}}"


def get_include_G_latex(report_data: TSPReportData):
    paths_way = report_data.paths_way
    worst_tau = report_data.worst_tau
    data = paths_way + [(worst_tau[0], True)]
    return f"G_{{ {paths_way_to_latex(data)} }}^{{ ({report_data.tree_level + 1})}}"


def get_previous_G_latex(report_data: TSPReportData | TSPReportDataLast) -> str:
    paths_way = report_data.paths_way
    data = paths_way
    under_latex = paths_way_to_latex(data)
    if under_latex == "()":
        return f"G^{{ ({report_data.tree_level})}}"
    else:
        return f"G_{{ {under_latex} }}^{{ ({report_data.tree_level})}}"


def matrix_to_table(city_matrix: TSPCityMatrix):
    result = []
    column_keys = None
    dict_matrix = city_matrix.matrix
    for i in sorted(dict_matrix.keys()):
        result.append([i + 1])
        if column_keys is None:
            column_keys = list(sorted(dict_matrix[i].keys()))
        for j in sorted(dict_matrix[i].keys()):
            if dict_matrix[i][j] > 10 ** 7:
                elem = latex2omml("\\infty")
            else:
                elem = latex2omml(f"{dict_matrix[i][j]}")
            result[-1].append(elem)
    result.insert(0, ["-"] + [i + 1 for i in column_keys])
    return result


class TSPStepDocxFiller(DocxFiller):
    def __init__(self, step_report_data: TSPReportData, step_number):
        self.report_data = step_report_data
        self.step_number = step_number

    def get_data_producers(self) -> Dict[str, Callable]:
        return {
            "step_number": self._produce_step_number,
            "current_G": self._produce_current_G,
            "start_table": self._produce_start_table,
            "taus": self._produce_taus,
            "u": self._produce_u,
            "exclude_G": self._produce_exclude_G,
            "exclude_v": self._produce_exclude_v,
            "table_removed_kl": self._produce_table_removed_kl,
            "table_filtered_paths": self._produce_table_filtered_paths,
            "table_zeroing": self._produce_table_zeroing,
            "h": self._produce_h,
            "include_G": self._produce_include_G,
            "include_v": self._produce_include_v,
            "result": self._produce_result
        }

    def _produce_step_number(self):
        return str(self.step_number + 1)

    def _produce_start_table(self):
        doc: DocumentType = Document()

        matrix = self.report_data.start_matrix
        tabulate_data = matrix_to_table(matrix)
        table = create_table_filled(document=doc, data=tabulate_data)
        paint_table_row(table, 0)
        paint_table_column(table, 0)
        return doc

    def _produce_taus(self):
        doc: DocumentType = Document()
        for tau in self.report_data.taus:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.element.append(latex2omml(
                f"\\tau\\left( {tau[0][0] + 1},{tau[0][1] + 1} \\right) = "
                f"{tau[2][0]} + {tau[2][1]} = {tau[1]} "))
        return doc

    def _produce_u(self):
        doc: DocumentType = Document()
        u = self.report_data.worst_tau
        doc.add_paragraph().add_run().element.append(
            latex2omml(f"u\\left( {u[0][0] + 1},{u[0][1] + 1} \\right) = {u[1]}"))
        return doc

    def _produce_exclude_G(self):
        doc: DocumentType = Document()
        doc.add_paragraph().add_run().element.append(latex2omml(get_exclude_G_latex(self.report_data)))
        return doc

    def _produce_exclude_v(self):
        doc: DocumentType = Document()
        worst_tau = self.report_data.worst_tau
        previous_v = self.report_data.previous_evaluation
        exclude_v = self.report_data.exclude_evaluation
        latex = (f"V({get_exclude_G_latex(self.report_data)}) = V({get_previous_G_latex(self.report_data)}) "
                 f"+ u\\left( {worst_tau[0][0] + 1},{worst_tau[0][1] + 1} \\right) = {previous_v} + {worst_tau[1]} = {exclude_v}")
        doc.add_paragraph().add_run().element.append(latex2omml(latex))

        return doc

    def _produce_table_removed_kl(self):
        trimmed_matrix = self.report_data.trimmed_matrix
        data = matrix_to_table(trimmed_matrix)
        doc: DocumentType = Document()
        table = create_table_filled(document=doc, data=data)
        paint_table_row(table, 0)
        paint_table_column(table, 0)
        return doc

    def _produce_table_filtered_paths(self):
        filtered_paths_matrix = self.report_data.filtered_paths_matrix
        data = matrix_to_table(filtered_paths_matrix)
        doc: DocumentType = Document()
        table = create_table_filled(document=doc, data=data)
        paint_table_row(table, 0)
        paint_table_column(table, 0)
        return doc

    def _produce_table_zeroing(self):
        zeroing_matrix = self.report_data.zeroing_matrix
        data = matrix_to_table(zeroing_matrix)
        doc: DocumentType = Document()
        table = create_table_filled(document=doc, data=data)
        paint_table_row(table, 0)
        paint_table_column(table, 0)
        return doc

    def _produce_h(self):
        doc: DocumentType = Document()
        terms = self.report_data.h_terms
        latex = f"h = {sum(terms)}"
        doc.add_paragraph().add_run().element.append(latex2omml(latex))
        return doc

    def _produce_include_G(self):
        doc: DocumentType = Document()
        doc.add_paragraph().add_run().element.append(latex2omml(get_include_G_latex(self.report_data)))
        return doc

    def _produce_include_v(self):
        doc: DocumentType = Document()
        terms = self.report_data.h_terms
        h = sum(terms)
        previous_evaluation = self.report_data.previous_evaluation
        include_evaluation = self.report_data.include_evaluation
        latex = (f"V({get_include_G_latex(self.report_data)}) = V({get_previous_G_latex(self.report_data)}) "
                 f"+ h = {previous_evaluation} + {h} = {include_evaluation}")
        doc.add_paragraph().add_run().element.append(latex2omml(latex))

        return doc

    def _produce_current_G(self):
        return latex2omml(get_previous_G_latex(self.report_data))

    def _produce_result(self):
        doc: DocumentType = Document()
        is_next_from_candidates = self.report_data.is_next_from_candidates
        p = doc.add_paragraph()
        if not is_next_from_candidates:
            try:
                is_included = self.report_data.is_included
                if is_included:
                    latex = f"V({get_include_G_latex(self.report_data)}) < V({get_exclude_G_latex(self.report_data)})"
                    text = " значит дальше идем по ветке, включающей вершину"
                else:
                    latex = f"V({get_exclude_G_latex(self.report_data)}) < V({get_include_G_latex(self.report_data)})"
                    text = (" значит дальше идем по ветке, исключающей вершину, кроме того преобразуем матрицу, как "
                            "указано в теории")
                p.add_run().element.append(latex2omml(
                    latex
                ))
                p.add_run(text)
            except:
                text = "Конец"
                p.add_run(text)
        else:
            text = ("В одной из висячих вершин оценочная функция V меньше чем любая на этом шаге. "
                    "Значит переходим к ней")
            p.add_run(text)
        return doc


class TSPLastStepDocxFiller(DocxFiller):
    def __init__(self, step_report_data: TSPReportDataLast, step_number):
        self.report_data = step_report_data
        self.step_number = step_number

    def get_data_producers(self) -> Dict[str, Callable]:
        return {
            "step_number": self._produce_step_number,
            "current_G": self._produce_current_G,
            "start_table": self._produce_start_table,
            "result_path": self._produce_result_path
        }

    def _produce_step_number(self):
        return str(self.step_number + 1)

    def _produce_current_G(self):
        return latex2omml(get_previous_G_latex(self.report_data))

    def _produce_start_table(self):
        doc: DocumentType = Document()

        matrix = self.report_data.start_matrix
        tabulate_data = matrix_to_table(matrix)
        table = create_table_filled(document=doc, data=tabulate_data)
        paint_table_row(table, 0)
        paint_table_column(table, 0)
        return doc

    def _produce_result_path(self):
        result_path = self.report_data.result_path

        graph = {}
        for edge in result_path:
            if edge[0] not in graph:
                graph[edge[0]] = edge[1]

        def dfs(node, _visited, segment):
            segment.append(node)
            if node in graph and not _visited.get(node, False):
                _visited[node] = True
                dfs(node=graph[node], _visited=_visited, segment=segment)

        visited = {node: False for node in graph}
        result = []
        dfs(node=result_path[0][0], _visited=visited, segment=result)

        latex = ' \\rightarrow '.join([str(i + 1) for i in result])
        return latex2omml(latex)


class TSPZeroStepDocxFiller(DocxFiller):
    def __init__(self, step_report_data: TSPReportDataMinimal):
        self.report_data = step_report_data

    def get_data_producers(self) -> Dict[str, Callable]:
        return {
            "table_zeroing": self._produce_table_zeroing,
            "h": lambda: latex2omml(f"{sum(self.report_data.h_terms)}"),
        }

    def _produce_table_zeroing(self):
        zeroing_matrix = self.report_data.zeroing_matrix
        data = matrix_to_table(zeroing_matrix)
        doc: DocumentType = Document()
        table = create_table_filled(document=doc, data=data)
        paint_table_row(table, 0)
        paint_table_column(table, 0)
        return doc
