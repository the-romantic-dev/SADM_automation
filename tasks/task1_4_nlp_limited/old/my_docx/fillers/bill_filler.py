from docx.shared import Pt
from sympy import symbols, Eq, latex, Matrix

from methods.bill.bill import Bill
from my_docx import docx_output
from my_docx.docx_output import latex2omml, sympy2omml
from task import Task
from util import merge_symplex_table_and_vars
from docx.document import Document as DocumentType
from docx import Document


class BillDocxFiller:
    def __init__(self, bill: Bill):
        self.bill: Bill = bill

    def _create_bill_lim1_canonical(self):
        x1, x2, x3 = symbols('x1 x2 x3')
        lim = Task.lim1234[0]
        const = -Task.lim1234[0].subs({x1: 0, x2: 0})
        return sympy2omml(Eq(lim + const + x3, const))

    def _create_bill_lim2_canonical(self):
        x1, x2, x4 = symbols('x1 x2 x4')
        lim = Task.lim1234[1]
        const = -Task.lim1234[1].subs({x1: 0, x2: 0})
        return sympy2omml(Eq(lim + const + x4, const))

    def _create_bill_tables(self):
        doc: DocumentType = Document()
        f = symbols('f')
        total_skips = 0
        for i in range(len(self.bill.tables_1)):
            base_latexs = [latex(var) for var in self.bill.all_base[i]]
            header_paragraph = doc.add_paragraph()
            header_run = header_paragraph.add_run('\nБазис ')
            header_run.font.size = Pt(14)
            header_run.element.append(latex2omml(' , '.join(base_latexs)))
            header_run.font.bold = True

            doc.add_paragraph().add_run().element.append(sympy2omml(Eq(f, self.bill.f_expressions[i])))
            data1 = merge_symplex_table_and_vars(
                symplex_table=self.bill.tables_1[i],
                free=self.bill.all_free[i],
                base=self.bill.all_base[i] + [r'\frac{∂f}{∂x_j},   \frac{∂f}{∂u_j}',
                                              r'\frac{∂f}{∂x_j}(X^* ),   \frac{∂f}{∂u_j}(X^* )']
            )
            table1 = docx_output.create_table_filled(
                document=doc,
                data=data1
            )
            table1.cell(0, 0).text = ''
            docx_output.paint_table_cell(table1, 0, 0, color='000000')
            docx_output.paint_table_cell(table1, -1, -1, color='000000')
            docx_output.paint_table_cell(table1, -2, -1, color='000000')

            if i < len(self.bill.tables_2):
                docx_output.paint_table_column(table1, self.bill.all_opt_cols[i] + 1)
                data2 = merge_symplex_table_and_vars(
                    symplex_table=self.bill.tables_2[i],
                    free=self.bill.all_free[i],
                    base=self.bill.all_base[i] + (
                        [f'u_{i + 1 - total_skips}'] if len(self.bill.tables_2[i]) > len(self.bill.all_base[i]) else [])
                )
                if self.bill.u_expressions[i] is None:
                    total_skips += 1
                    doc.add_paragraph().add_run()
                else:
                    doc.add_paragraph().add_run().element.append(
                        sympy2omml(Eq(symbols(f'u{i + 1 - total_skips}'), self.bill.u_expressions[i])))

                table2 = docx_output.create_table_filled(
                    document=doc,
                    data=data2
                )
                table2.cell(0, 0).text = ''
                docx_output.paint_table_cell(table2, 0, 0, color='000000')
                docx_output.paint_table_column(table2, self.bill.all_opt_cols[i] + 1)
                docx_output.paint_table_row(table2, self.bill.all_opt_rows[i] + 1)
        return doc

    def _create_bill_pic(self):
        return 'report_data/bill.png'

    def _crate_bill_X(self):
        return latex2omml('X^* = ' + latex(Matrix([self.bill.solution['x1'], self.bill.solution['x2']])) + r'\approx' +
                          latex(Matrix(
                              [round(float(self.bill.solution['x1']), 4), round(float(self.bill.solution['x2']), 4)])))

    def _create_bill_f(self):
        return latex2omml('f = ' + latex(self.bill.solution['f']) + r'\approx' + latex(round(float(self.bill.solution['f']), 4)))
    def get_data_producers(self):
        return {
            'lim1_canonical': self._create_bill_lim1_canonical,
            'lim2_canonical': self._create_bill_lim2_canonical,
            'bill_tables': self._create_bill_tables,
            'bill_pic': self._create_bill_pic,
            'bill_X': self._crate_bill_X,
            'bill_f': self._create_bill_f
        }
