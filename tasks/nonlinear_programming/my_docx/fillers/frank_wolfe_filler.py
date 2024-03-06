from docx import Document
from docx.document import Document as DocumentType
from sympy import Matrix, eye, zeros, latex

from methods.frank_wolfe.frank_wolfe import FrankWolfe
from my_docx import docx_output
from my_docx.docx_output import sympy2omml, latex2omml
from util import merge_symplex_table_and_vars


class FrankWolfeDocxFiller:
    def __init__(self, frank_wolfe: FrankWolfe):
        self.frank_wolfe: FrankWolfe = frank_wolfe

    def get_data_producers(self):
        return {
            'fw_p': lambda: sympy2omml(self.frank_wolfe.p),
            'fw_b': lambda: sympy2omml(self.frank_wolfe.b),
            'fw_H': lambda: sympy2omml(self.frank_wolfe.H),
            'fw_A': lambda: sympy2omml(self.frank_wolfe.A),
            'fw_AEOO_HOEAT': lambda: sympy2omml(self.frank_wolfe.AEOO_HOEAT),
            'fw_bp': lambda: sympy2omml(self.frank_wolfe.bp),
            'fw_AEOO_HOEAT_minus': self._create_AEOO_HOEAT_minus,
            'fw_bp_minus': self._create_bp_minus,
            'fw_kk_AEOO_HOEAT': self._create_kk_AEOO_HOEAT,
            'fw_kk_tables': self._create_kk_tables,
            'fw_tables': self._create_fw_tables,
            'fw_pic': self._create_fw_pic,
            'fw_X': self._crate_fw_X,
            'fw_f': self._create_fw_f
        }

    def _create_AEOO_HOEAT_minus(self):
        AEOO_HOEAT: Matrix = self.frank_wolfe.AEOO_HOEAT.copy()
        AEOO_HOEAT.row_op(2, lambda v, j: -v)
        AEOO_HOEAT.row_op(3, lambda v, j: -v)
        return sympy2omml(AEOO_HOEAT)

    def _create_bp_minus(self):
        bp: Matrix = self.frank_wolfe.bp.copy()
        bp.row_op(2, lambda v, j: -v)
        bp.row_op(3, lambda v, j: -v)
        return sympy2omml(bp)

    def _create_kk_AEOO_HOEAT(self):
        AEOO_HOEAT: Matrix = self.frank_wolfe.AEOO_HOEAT.copy()
        AEOO_HOEAT.row_op(2, lambda v, j: -v)
        AEOO_HOEAT.row_op(3, lambda v, j: -v)

        E = eye(2)
        O = zeros(2)
        K = O.col_join(E)

        return sympy2omml(AEOO_HOEAT.row_join(K))

    def _create_kk_tables(self):
        doc: DocumentType = Document()
        for i in range(len(self.frank_wolfe.kk_symplex_track)):
            data = merge_symplex_table_and_vars(
                symplex_table=self.frank_wolfe.kk_symplex_track[i],
                free=self.frank_wolfe.kk_free_track[i],
                base=self.frank_wolfe.kk_base_track[i] + ['- k_1 - k_2']
            )
            table = docx_output.create_table_filled(
                document=doc,
                data=data
            )
            doc.add_paragraph()
            docx_output.paint_table_cell(table, 0, 0, color='000000')
            table.cell(0, 0).text = ""
            if i < len(self.frank_wolfe.kk_symplex_track) - 1:
                docx_output.paint_table_column(table, self.frank_wolfe.kk_opt_col_track[i] + 1)
                docx_output.paint_table_row(table, self.frank_wolfe.kk_opt_row_track[i] + 1)
        return doc

    def _create_fw_tables(self):
        doc: DocumentType = Document()
        tables_track = self.frank_wolfe.fw_symplex_track
        free_track = self.frank_wolfe.fw_free_track
        base_track = self.frank_wolfe.fw_base_track
        old_i = -1
        last_j = -1
        for i in range(len(tables_track)):
            for j in range(len(tables_track[i])):
                if old_i != i:
                    old_i = i

                    if i != 0:
                        Z_ji = r'Z_{ji}'.replace('{ji}', '{' + str(last_j + 1) + str(i-1) + '}')
                        Z_i =  r'Z_{i}'.replace('{i}', '{' + str(i-1) + '}')
                        tilda_Z_i = r'\widetilde{Z}_{i}'.replace('{i}', '{' + str(i-1) + '}')
                        tilda_Z_ji = r'\widetilde{Z}_{ji}'.replace('{ji}', '{' + str(last_j + 1) + str(i-1) + '}')
                        doc.add_paragraph().add_run("Так как значение функции уменьшилось более чем в 2 раза, находим новое Z")
                        doc.add_paragraph().add_run().element.append(latex2omml(r't=\min\left(1,t^* \right)'))
                        doc.add_paragraph().add_run().element.append(latex2omml(
                            (r't^*=-\frac{\left( Z_{ji} - Z_{i} \right)^T \widetilde{Z}_{i}}' +
                            r'{\left( Z_{ji} - Z_{i} \right)^T\left( \widetilde{Z}_{ji} - \widetilde{Z}_{i} \right)}').
                            replace('{i}', '{' + str(i-1) + '}').
                            replace('{ji}', '{' + str(last_j + 1) + str(i-1) + '}') + f' = {latex(self.frank_wolfe.t_default_track[i-1])}'
                        ))
                        doc.add_paragraph().add_run().element.append(latex2omml(f't = {self.frank_wolfe.t_track[i-1]}'))
                        doc.add_paragraph().add_run().element.append(latex2omml(f'Z_{i} = {Z_i} + t({Z_ji} - {Z_i})'))
                    Z_expr_latex = (r'Z_{ji} = '.
                                    replace('i', str(i)).
                                    replace('j', str(j + 1)) +
                                    f'{latex(self.frank_wolfe.Z_expr_track[i])}')
                    tilda_Z_latex = (r'\widetilde{Z}_{i} = '
                                     .replace('{i}', '{' + str(i) + '}') +
                                     f'{latex(self.frank_wolfe.tilda_Z_track[i])}')
                    doc.add_paragraph().add_run().element.append(latex2omml(Z_expr_latex + ' , ' + tilda_Z_latex))
                table = tables_track[i][j]
                free = free_track[i][j]
                base = base_track[i][j]
                data = merge_symplex_table_and_vars(
                    symplex_table=table,
                    free=free,
                    base=base + ['Z_{ji}\widetilde{Z}_{i}'.
                                 replace('{i}', '{' + str(i) + '}').
                                 replace('{ji}', '{' + str(j + 1) + str(i) + '}')]
                )
                table = docx_output.create_table_filled(
                    document=doc,
                    data=data
                )
                docx_output.paint_table_cell(table, 0, 0, color='000000')
                table.cell(0, 0).text = ""
                if j < len(self.frank_wolfe.fw_symplex_track[i]) - 1:
                    docx_output.paint_table_column(table, self.frank_wolfe.fw_opt_col_track[i][j] + 1)
                    docx_output.paint_table_row(table, self.frank_wolfe.fw_opt_row_track[i][j] + 1)
                doc.add_paragraph()
                last_j = j
        return doc

    def _create_fw_pic(self):
        return 'report_data/frank_wolfe.png'

    def _crate_fw_X(self):
        return latex2omml('X^* = ' + latex(Matrix([self.frank_wolfe.solution['x1'], self.frank_wolfe.solution['x2']])) + r'\approx' +
                          latex(Matrix(
                              [round(float(self.frank_wolfe.solution['x1']), 4), round(float(self.frank_wolfe.solution['x2']), 4)])))

    def _create_fw_f(self):
        return latex2omml('f = ' + latex(self.frank_wolfe.solution['f']) + r'\approx' + latex(round(float(self.frank_wolfe.solution['f']), 4)))